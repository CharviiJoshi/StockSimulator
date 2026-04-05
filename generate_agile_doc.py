import os, sys, subprocess
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# STANDARD A4/LETTER PORTRAIT with 0.75 inch margins (usable width = 7 inches)
for s in doc.sections:
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.underline = True; r.font.size = Pt(16)

def subtitle(text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)

def make_table(headers, data, col_widths, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False; t.allow_autofit = False
    for i, w in enumerate(col_widths):
        for c in t.columns[i].cells: c.width = w
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(font_size)
        t.rows[0].cells[i].paragraphs[0].paragraph_format.space_before = Pt(3)
        t.rows[0].cells[i].paragraphs[0].paragraph_format.space_after = Pt(3)
    for rd in data:
        row = t.add_row().cells
        for i in range(len(headers)):
            row[i].text = str(rd[i])
            for run in row[i].paragraphs[0].runs:
                run.font.size = Pt(font_size)
            row[i].paragraphs[0].paragraph_format.space_before = Pt(4)
            row[i].paragraphs[0].paragraph_format.space_after = Pt(4)

# ===================== 1. SPRINT RETROSPECTIVE =====================
title('SPRINT RETROSPECTIVE')
doc.add_paragraph()

# 3-col table: 1.2 + 2.9 + 2.9 = 7.0 inches
make_table(
    ['Sprint Number', 'What Went Well', 'What Went Wrong'],
    [
        ("Sprint 1\n(Planning)\n16/02–22/02",
         "• Finalized SRS mapped to core modules.\n• Delivered UML Class, Sequence, Activity diagrams.\n• Strict NFRs defined for FastAPI.\n• GitHub repo created with CI/CD.",
         "• Ambiguity mapping objects to SQLite.\n• Overlapping actor definitions in UML.\n• SRS formatting to IEEE took extra time.\n• Memento state scope was unclear."),
        ("Sprint 2\n(Auth & UI)\n23/02–01/03",
         "• Full Auth flow (Login, Register, Profile).\n• React integrated with FastAPI backend.\n• SMTP mail service for Forgot Password.\n• Reusable Header/Footer components.",
         "• Prop-drilling before state management.\n• CORS errors blocked API testing.\n• Async DB connections dropped requests.\n• JWT refresh had expiration bugs."),
        ("Sprint 3\n(Core Engine)\n02/03–08/03",
         "• Day-by-day simulation loop built.\n• Portfolio with Buy/Sell logic integrated.\n• DB transaction integrity verified.\n• Memento undo feature implemented.",
         "• Large CSVs caused memory overflows.\n• Date gaps required schema changes.\n• Atomic transactions were complex.\n• Timer hooks caused state desyncs."),
        ("Sprint 4\n(Dashboards)\n09/03–15/03",
         "• Net-Worth Dashboard with Pie Charts.\n• Gamification (Streaks, Badges) added.\n• PDF generation for monthly reports.\n• E2E tests passed before sign-off.",
         "• Pagination degraded History loads.\n• Chart init lagged on large datasets.\n• Cache vs Undo state consistency bugs.\n• PDF page-breaks needed iterations.")
    ],
    [Inches(1.2), Inches(2.9), Inches(2.9)], font_size=10)

doc.add_page_break()

# ===================== 2. SPRINT REVIEW =====================
title('SPRINT REVIEW')

rh = ['Day', 'What went well', "What didn't go well", 'What still Puzzles us', 'What Have we Learnt']
# 5-col: 1.0 + 1.5*4 = 7.0 inches
rw = [Inches(1.0), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5)]

subtitle('SPRINT 1 - REVIEW (16/02 – 22/02)')
make_table(rh, [
    ("Day 1\n16/02", "Defined architecture scope.", "Actor role ambiguity.", "DB schema complexity?", "Keep early models simple."),
    ("Day 2\n17/02", "Refined primary use cases.", "Use-case flow overlaps.", "Best granularity level?", "Merge similar components."),
    ("Day 3\n18/02", "Documented core requirements.", "NFRs took extra time.", "Measuring performance?", "Quantifiable metrics key."),
    ("Day 4\n19/02", "Finalized Class and Sequence.", "UML formatting delays.", "Handling async flows?", "Collaborative tools help."),
    ("Day 5-7\n20–22/02", "SRS updates and reviews.", "Feedback conflicts arose.", "Covering edge cases?", "Iterative reviews help.")
], rw, font_size=10)

subtitle('SPRINT 2 - REVIEW (23/02 – 01/03)')
make_table(rh, [
    ("Day 1\n23/02", "Created Header and Homepage.", "Responsive layout issues.", "Best breakpoints?", "Modular approach saves time."),
    ("Day 2\n24/02", "Login/Register forms built.", "Multi-step state issues.", "Async validation edges?", "Centralized logic is clean."),
    ("Day 3\n25/02", "Forgot Password with SMTP.", "Secure token pipelines.", "Optimal token expiry?", "Short-lived tokens critical."),
    ("Day 4\n26/02", "UI acceptance criteria done.", "Detail vs brevity balance.", "Criteria precision level?", "BDD formats clarify scope."),
    ("Day 5-7\n27/02–\n01/03", "Auth linked to SQLite DB.", "Async DB connection drops.", "API error frameworks?", "Global middleware stabilizes.")
], rw, font_size=10)

subtitle('SPRINT 3 - REVIEW (02/03 – 08/03)')
make_table(rh, [
    ("Day 1\n02/03", "CSV parsing implemented.", "Memory on large CSVs.", "Stream vs chunk load?", "Streaming wins entirely."),
    ("Day 2\n03/03", "Portfolio state initialized.", "Schema design for history.", "Optimal snapshot index?", "Normalization is faster."),
    ("Day 3\n04/03", "Simulation loop working.", "Loop performance issues.", "Non-blocking fetches?", "Batch ops reduce overhead."),
    ("Day 4\n05/03", "Buy/Sell logic completed.", "Insufficient fund errors.", "Concurrent DB safety?", "Atomic transactions needed."),
    ("Day 5-7\n06–08/03", "Memento undo integrated.", "Cascading state revert.", "Event vs snapshot undo?", "Snapshots simpler to use.")
], rw, font_size=10)

subtitle('SPRINT 4 - REVIEW (09/03 – 15/03)')
make_table(rh, [
    ("Day 1\n09/03", "Trading history table built.", "Pagination was very slow.", "Better indexing method?", "Keyset pagination scales."),
    ("Day 2\n10/03", "Net-Worth charts added.", "Canvas rendering lagged.", "Custom chart tooltips?", "Server-side aggregation."),
    ("Day 3\n11/03", "Gamification elements added.", "Streak rule definition.", "Prevent gaming abuse?", "Backend validation is key."),
    ("Day 4\n12/03", "Report aggregation complete.", "PDF export formatting.", "Best PDF generator?", "Dedicated libs work best."),
    ("Day 5-7\n13–15/03", "E2E integration finished.", "Minor frontend desyncs.", "Full E2E automation?", "CI/CD checks are reliable.")
], rw, font_size=10)

doc.add_page_break()

# ===================== 3. DAILY STANDUP MEETINGS =====================
title('DAILY STANDUP MEETINGS')

sh = ['Day & Date', 'Team Members', 'What you did Yesterday?', 'What you will do Today?', 'Obstacles faced']
# 5-col: 0.9 + 1.3 + 1.6*3 = 7.0
sw = [Inches(0.9), Inches(1.3), Inches(1.6), Inches(1.6), Inches(1.6)]

T = ["Ayeshaani Kushwaha", "Charvi Joshi", "Ayush Solanki", "Chirag Hoondlani", "Dherya Mujawdiya"]

# SPRINT 1
s1 = [
    ("Day 1\n16/02", T[0], "Reviewed project requirements.", "Starting Use-Case Diagram.", "Understanding actor roles."),
    ("Day 1\n16/02", T[1], "Analyzed system workflow.", "Created GitHub repository.", "Setting up access rules."),
    ("Day 1\n16/02", T[2], "Studied DB design patterns.", "Starting Activity Diagram.", "Clarifying process order."),
    ("Day 1\n16/02", T[3], "Identified entity relations.", "Drafting database schema.", "Deciding FK constraints."),
    ("Day 1\n16/02", T[4], "Reviewed SRS templates.", "Starting UI wireframes.", "Understanding UI scope."),
    ("Day 2\n17/02", T[0], "Completed Use-Case draft.", "Refining sub-flows.", "Overlapping actor duties."),
    ("Day 2\n17/02", T[1], "Finished GitHub repo setup.", "Designing Class Diagram.", "Mapping class attributes."),
    ("Day 2\n17/02", T[2], "Started Activity Diagram.", "Completing all paths.", "Finalizing flow paths."),
    ("Day 2\n17/02", T[3], "Drafted baseline schema.", "Refining DB relations.", "Composite key scenarios."),
    ("Day 2\n17/02", T[4], "Started dashboard wireframes.", "Designing main layout.", "Layout alignment issues."),
    ("Day 3\n18/02", T[0], "Refined Use-Case sub-flows.", "Cross-ref with Sequence.", "Syncing with Class Diagram."),
    ("Day 3\n18/02", T[1], "Completed Class Diagram.", "Designing Sequence Diagram.", "Clarifying method calls."),
    ("Day 3\n18/02", T[2], "Completed Activity Diagram.", "Documenting functional reqs.", "IEEE formatting rules."),
    ("Day 3\n18/02", T[3], "Refined entity relations.", "Drafting NFRs in SRS.", "Defining perf limits."),
    ("Day 3\n18/02", T[4], "Designed dashboard layout.", "Designing Auth page UI.", "Consistent styling."),
    ("Day 4\n19/02", T[0], "Helped finalize Sequence.", "Reviewing all diagrams.", "Standard UML notation."),
    ("Day 4\n19/02", T[1], "Completed Sequence Diagram.", "Verifying diagram consistency.", "Fixing naming conflicts."),
    ("Day 4\n19/02", T[2], "Documented functional reqs.", "Consolidating full SRS.", "Spacing/format conflicts."),
    ("Day 4\n19/02", T[3], "Drafted complete NFRs.", "Reviewing NFRs with team.", "Aligning with tech stack."),
    ("Day 4\n19/02", T[4], "Designed Auth page UI.", "Reviewing wireframes.", "Edge case view states."),
    ("Day 5\n20/02", T[0], "Reviewed all UML diagrams.", "Planning UI components.", "Categorizing reusable parts."),
    ("Day 5\n20/02", T[1], "Verified diagram consistency.", "Finalizing CI/CD pipeline.", "GitHub Actions config."),
    ("Day 5\n20/02", T[2], "Consolidated SRS document.", "Final SRS proofreading.", "Catching minor typos."),
    ("Day 5\n20/02", T[3], "Reviewed NFRs with team.", "Planning backend routes.", "REST payload structure."),
    ("Day 5\n20/02", T[4], "Reviewed all wireframes.", "Setting up React + Vite.", "Build config errors."),
    ("Day 6\n21/02", T[0], "Planned UI component tree.", "Finalizing component naming.", "Avoiding naming conflicts."),
    ("Day 6\n21/02", T[1], "Configured GitHub Actions.", "Testing CI build pipeline.", "Build script path errors."),
    ("Day 6\n21/02", T[2], "Proofread SRS functional reqs.", "Formatting SRS appendices.", "Table formatting issues."),
    ("Day 6\n21/02", T[3], "Planned backend route paths.", "Defining API response models.", "Choosing schema validators."),
    ("Day 6\n21/02", T[4], "Set up React + Vite project.", "Installing core dependencies.", "Dependency version clashes."),
    ("Day 7\n22/02", T[0], "Finalized component hierarchy.", "Sprint 1 retrospective prep.", "Summarizing key learnings."),
    ("Day 7\n22/02", T[1], "Tested CI/CD pipeline fully.", "Sprint 1 retrospective prep.", "Documenting blockers faced."),
    ("Day 7\n22/02", T[2], "Formatted SRS appendices.", "Sprint 1 retrospective prep.", "Aligning sections with IEEE."),
    ("Day 7\n22/02", T[3], "Defined API response schemas.", "Sprint 1 retrospective prep.", "Finalizing route naming."),
    ("Day 7\n22/02", T[4], "Installed all dependencies.", "Sprint 1 retrospective prep.", "Verifying dev environment."),
]

# SPRINT 2
s2 = [
    ("Day 1\n23/02", T[0], "Planned component hierarchy.", "Creating Header/Footer.", "Responsive design issues."),
    ("Day 1\n23/02", T[1], "Finalized CI/CD pipeline.", "Building backend + Login.", "Connecting SQLite securely."),
    ("Day 1\n23/02", T[2], "Concluded SRS verification.", "Designing BDD criteria.", "Standardizing test format."),
    ("Day 1\n23/02", T[3], "Planned REST endpoints.", "Implementing DB models.", "Async init timeout issues."),
    ("Day 1\n23/02", T[4], "Finished React env setup.", "Creating route guards.", "Protecting private routes."),
    ("Day 2\n24/02", T[0], "Created Header/Footer.", "Developing Registration UI.", "Frontend form validation."),
    ("Day 2\n24/02", T[1], "Built Login backend logic.", "Linked Auth to database.", "Password hashing issues."),
    ("Day 2\n24/02", T[2], "Created BDD criteria doc.", "Writing component tests.", "Mapping test identifiers."),
    ("Day 2\n24/02", T[3], "Implemented DB model layer.", "Creating Auth endpoints.", "JWT token generation."),
    ("Day 2\n24/02", T[4], "Created router with guards.", "Developing Homepage UI.", "Dynamic child components."),
    ("Day 3\n25/02", T[0], "Completed Registration UI.", "Developing Login Page UI.", "Prop-drilling auth state."),
    ("Day 3\n25/02", T[1], "Linked Auth to database.", "Building Forgot Pwd backend.", "Secure reset token hashes."),
    ("Day 3\n25/02", T[2], "Implemented component tests.", "Writing integration tests.", "Mocking DB responses."),
    ("Day 3\n25/02", T[3], "Created Auth endpoints.", "Adding rate limiting.", "Request timeout values."),
    ("Day 3\n25/02", T[4], "Developed Homepage UI.", "Building Edit Profile UI.", "Isolated form states."),
    ("Day 4\n26/02", T[0], "Finalized Login Page UI.", "Creating Forgot Pwd UI.", "OTP verification input."),
    ("Day 4\n26/02", T[1], "Built Forgot Pwd backend.", "Integrated SMTP mail service.", "SMTP relay config."),
    ("Day 4\n26/02", T[2], "Compiled test suite.", "Executing BDD checklists.", "Missing edge validations."),
    ("Day 4\n26/02", T[3], "Refined endpoints + limits.", "Linking APIs to frontend.", "CORS policy errors."),
    ("Day 4\n26/02", T[4], "Developed Profile Edit UI.", "Refactoring Redux slice.", "Session persistence."),
    ("Day 5\n27/02", T[0], "Polished Forgot Pwd UI.", "Reviewing all Auth UI.", "Toast notification styles."),
    ("Day 5\n27/02", T[1], "Tested SMTP mail service.", "Verified full Auth + DB flow.", "Minor token timing bugs."),
    ("Day 5\n27/02", T[2], "Executed BDD checklists.", "Documenting test results.", "Edge case summaries."),
    ("Day 5\n27/02", T[3], "Connected APIs to frontend.", "Fixing CORS integration.", "JSON error format."),
    ("Day 5\n27/02", T[4], "Refactored Redux auth.", "Preparing simulation setup.", "Nested routing structure."),
    ("Day 6\n28/02", T[0], "Reviewed all Auth UI flows.", "Cross-browser testing Auth.", "Safari input rendering bugs."),
    ("Day 6\n28/02", T[1], "Verified Auth + DB fully.", "Testing edge case auth flows.", "Empty field submit handling."),
    ("Day 6\n28/02", T[2], "Documented test results.", "Writing regression test plan.", "Prioritizing critical tests."),
    ("Day 6\n28/02", T[3], "Fixed CORS fully.", "Reviewing API error handling.", "Standardizing HTTP codes."),
    ("Day 6\n28/02", T[4], "Prepared simulation dirs.", "Planning simulation state mgmt.", "Redux vs Context decision."),
    ("Day 7\n01/03", T[0], "Completed cross-browser tests.", "Sprint 2 retrospective prep.", "Documenting UI fixes."),
    ("Day 7\n01/03", T[1], "Tested all auth edge cases.", "Sprint 2 retrospective prep.", "Summarizing backend status."),
    ("Day 7\n01/03", T[2], "Wrote regression test plan.", "Sprint 2 retrospective prep.", "Compiling test coverage."),
    ("Day 7\n01/03", T[3], "Reviewed all error handling.", "Sprint 2 retrospective prep.", "Listing open API issues."),
    ("Day 7\n01/03", T[4], "Planned simulation state.", "Sprint 2 retrospective prep.", "Summarizing frontend status."),
]

# SPRINT 3
s3 = [
    ("Day 1\n02/03", T[0], "Reviewed all auth UI flows.", "Designing CSV upload UI.", "Large file drag-and-drop."),
    ("Day 1\n02/03", T[1], "Verified auth DB integrity.", "Building CSV parser backend.", "Memory on large CSVs."),
    ("Day 1\n02/03", T[2], "Documented Sprint 2 tests.", "Designing data structures.", "Daily tick model size."),
    ("Day 1\n02/03", T[3], "Fixed CORS and JSON issues.", "Developing Portfolio Init API.", "User-session linking."),
    ("Day 1\n02/03", T[4], "Prepared simulation routing.", "Planning frontend loop.", "UI re-render frequency."),
    ("Day 2\n03/03", T[0], "Designed CSV Upload with preview.", "Integrating CSV endpoint.", "Multipart form parsing."),
    ("Day 2\n03/03", T[1], "Built Market Data parser.", "Verified CSV ingestion.", "Inconsistent date formats."),
    ("Day 2\n03/03", T[2], "Designed data structure schemas.", "Implementing Portfolio logic.", "Portfolio snapshot schema."),
    ("Day 2\n03/03", T[3], "Developed Portfolio Init API.", "Creating simulation loop.", "Sequential DB I/O perf."),
    ("Day 2\n03/03", T[4], "Planned frontend loop render.", "Connecting Play/Pause controls.", "Timer hook race conditions."),
    ("Day 3\n04/03", T[0], "Integrated CSV to backend.", "Refining Market Data view.", "Virtual scrolling for data."),
    ("Day 3\n04/03", T[1], "Verified ingestion quality.", "Added market data validation.", "Missing trading day gaps."),
    ("Day 3\n04/03", T[2], "Implemented portfolio tracking.", "Writing Buy algorithm.", "Insufficient funds case."),
    ("Day 3\n04/03", T[3], "Created backend sim loop.", "Optimizing DB fetch in loop.", "DB I/O blocking thread."),
    ("Day 3\n04/03", T[4], "Connected Play/Pause controls.", "Implementing Sell algorithm.", "Avg cost basis updates."),
    ("Day 4\n05/03", T[0], "Refined Market Data view.", "Preparing Trading UI forms.", "Stock quantity validation."),
    ("Day 4\n05/03", T[1], "Added full data validation.", "Verified Buy/Sell end-to-end.", "Atomic DB transactions."),
    ("Day 4\n05/03", T[2], "Wrote Buy algorithm logic.", "Reviewing response contracts.", "Transaction ledger format."),
    ("Day 4\n05/03", T[3], "Optimized DB in sim loop.", "Integrating Undo (Memento).", "State recreation on undo."),
    ("Day 4\n05/03", T[4], "Implemented Sell algorithm.", "Connecting Trading UI to API.", "Debouncing rapid trades."),
    ("Day 5\n06/03", T[0], "Polished Trading UI forms.", "Testing full trading flow.", "Minor state desyncs."),
    ("Day 5\n06/03", T[1], "Verified Buy/Sell interactions.", "Ensured backend consistency.", "Rollback edge bugs."),
    ("Day 5\n06/03", T[2], "Reviewed response structures.", "Writing Undo unit tests.", "Mocking state chains."),
    ("Day 5\n06/03", T[3], "Integrated Memento undo.", "Refining state snapshots.", "Memory overhead of states."),
    ("Day 5\n06/03", T[4], "Connected Trading UI to APIs.", "Reviewing simulation UX.", "Trade feedback toasts."),
    ("Day 6\n07/03", T[0], "Tested full trading flow.", "Fixing UI state edge cases.", "Stale data after undo."),
    ("Day 6\n07/03", T[1], "Ensured backend consistency.", "Stress testing simulation loop.", "High-load DB timeout errors."),
    ("Day 6\n07/03", T[2], "Wrote Undo unit tests.", "Running full test suite.", "Flaky async test results."),
    ("Day 6\n07/03", T[3], "Refined state snapshot logic.", "Optimizing snapshot storage.", "Pruning old snapshot data."),
    ("Day 6\n07/03", T[4], "Reviewed simulation UX flow.", "Adding loading indicators.", "Spinner placement issues."),
    ("Day 7\n08/03", T[0], "Fixed UI state edge cases.", "Sprint 3 retrospective prep.", "Documenting UI bug fixes."),
    ("Day 7\n08/03", T[1], "Stress tested simulation loop.", "Sprint 3 retrospective prep.", "Summarizing perf results."),
    ("Day 7\n08/03", T[2], "Ran full test suite.", "Sprint 3 retrospective prep.", "Compiling test pass rate."),
    ("Day 7\n08/03", T[3], "Optimized snapshot storage.", "Sprint 3 retrospective prep.", "Documenting Memento changes."),
    ("Day 7\n08/03", T[4], "Added loading indicators.", "Sprint 3 retrospective prep.", "Summarizing UX improvements."),
]

# SPRINT 4
s4 = [
    ("Day 1\n09/03", T[0], "Tested trading UI thoroughly.", "Designing dashboard charts.", "Choosing chart library."),
    ("Day 1\n09/03", T[1], "Ensured backend consistency.", "Writing Net-Worth APIs.", "Aggregating large datasets."),
    ("Day 1\n09/03", T[2], "Wrote Undo feature tests.", "Implementing Pie Chart UI.", "Dynamic data to chart."),
    ("Day 1\n09/03", T[3], "Refined state snapshotting.", "Developing History view.", "Pagination performance."),
    ("Day 1\n09/03", T[4], "Reviewed simulation UX flow.", "Implementing Streak rules.", "Streak conditions logic."),
    ("Day 2\n10/03", T[0], "Designed dashboard layout.", "Integrating charts with API.", "Chart rendering lag."),
    ("Day 2\n10/03", T[1], "Wrote Net-Worth API routes.", "Verified Dashboard data flow.", "Time-series formatting."),
    ("Day 2\n10/03", T[2], "Implemented Pie Chart view.", "Implementing Badge logic.", "Badge threshold values."),
    ("Day 2\n10/03", T[3], "Developed History list view.", "Optimizing History queries.", "Adding DB indexes."),
    ("Day 2\n10/03", T[4], "Implemented Streak backend.", "Linking Streak to frontend.", "Instant UI updates."),
    ("Day 3\n11/03", T[0], "Integrated charts with API.", "Styling badge UI elements.", "Badge visual design."),
    ("Day 3\n11/03", T[1], "Verified dashboard integrity.", "Setting up Report aggregation.", "Complex SQL aggregation."),
    ("Day 3\n11/03", T[2], "Implemented Badge system.", "Connecting Badges to UI.", "Multi-context state sync."),
    ("Day 3\n11/03", T[3], "Optimized queries + indexes.", "Developing PDF generation.", "Headless PDF formatting."),
    ("Day 3\n11/03", T[4], "Linked Streak to frontend.", "Generating test dummy data.", "Test data edge cases."),
    ("Day 4\n12/03", T[0], "Styled Gamification UI.", "Reviewing cross-browser UI.", "CSS overflow on mobile."),
    ("Day 4\n12/03", T[1], "Set up Report aggregation.", "Verified Reports engine.", "Rounding calculation bugs."),
    ("Day 4\n12/03", T[2], "Connected Badges API to UI.", "Refining PDF styling.", "PDF page break markers."),
    ("Day 4\n12/03", T[3], "Developed PDF gen service.", "Running E2E tests.", "Isolating flaky tests."),
    ("Day 4\n12/03", T[4], "Generated test QA data.", "Assisting E2E automation.", "Simulating user journeys."),
    ("Day 5\n13/03", T[0], "Reviewed full UI layout.", "Finalizing frontend build.", "Minified bundle issues."),
    ("Day 5\n13/03", T[1], "Verified Reports accuracy.", "Final review + deploy prep.", "Env variable mapping."),
    ("Day 5\n13/03", T[2], "Refined PDF page layout.", "Executing final QA checks.", "No major issues found."),
    ("Day 5\n13/03", T[3], "Ran full E2E test suites.", "Resolving final UI bugs.", "Timezone date shift bugs."),
    ("Day 5\n13/03", T[4], "Assisted E2E automation.", "Gathering sprint metrics.", "Compiling metric data."),
    ("Day 6\n14/03", T[0], "Finalized frontend build.", "Deploying to staging server.", "Build env discrepancies."),
    ("Day 6\n14/03", T[1], "Completed deploy prep.", "Running staging smoke tests.", "Staging DB config mismatch."),
    ("Day 6\n14/03", T[2], "Executed final QA checks.", "Documenting QA sign-off report.", "Formatting final report."),
    ("Day 6\n14/03", T[3], "Resolved final UI bugs.", "Verifying staging deployment.", "Network proxy timeouts."),
    ("Day 6\n14/03", T[4], "Gathered sprint metrics.", "Preparing demo presentation.", "Selecting key demo flows."),
    ("Day 7\n15/03", T[0], "Deployed to staging server.", "Sprint 4 retrospective prep.", "Documenting final UI state."),
    ("Day 7\n15/03", T[1], "Ran staging smoke tests.", "Sprint 4 retrospective prep.", "Summarizing deployment status."),
    ("Day 7\n15/03", T[2], "Documented QA sign-off.", "Sprint 4 retrospective prep.", "Compiling final test metrics."),
    ("Day 7\n15/03", T[3], "Verified staging deployment.", "Sprint 4 retrospective prep.", "Listing known minor issues."),
    ("Day 7\n15/03", T[4], "Prepared demo presentation.", "Sprint 4 retrospective prep.", "Finalizing project summary."),
]

subtitle('Sprint 1:- (16/02/26 – 22/02/26)')
make_table(sh, s1, sw, font_size=10)
subtitle('Sprint 2:- (23/02/26 – 01/03/26)')
make_table(sh, s2, sw, font_size=10)
subtitle('Sprint 3:- (02/03/26 – 08/03/26)')
make_table(sh, s3, sw, font_size=10)
subtitle('Sprint 4:- (09/03/26 – 15/03/26)')
make_table(sh, s4, sw, font_size=10)

os.makedirs('documentation', exist_ok=True)
fp = os.path.join('documentation', 'Agile_Documentation.docx')
doc.save(fp)
print(f"Success: Generated {fp}")
